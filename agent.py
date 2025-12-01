
import os, io, shutil, uuid, sys, subprocess
from dataclasses import dataclass
from typing import List

# Install dependencies if missing
def ensure(pkg: str):
    try:
        __import__(pkg.split("==")[0].replace("-", "_"))
    except Exception:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "-q", pkg])

for p in [
    "python-dotenv", "gradio", "PyPDF2", "python-docx",
    "qdrant-client", "langchain-text-splitters", "langchain-huggingface",
    "langgraph", "portkey-ai", "langchain-core"
]:
    ensure(p)

import gradio as gr
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document
from qdrant_client import QdrantClient
from qdrant_client.http import models as qmodels
from langgraph.graph import StateGraph, END
from portkey_ai import Portkey

# Config
load_dotenv()
EMBED_MODEL_NAME = "sentence-transformers/all-mpnet-base-v2"
DEFAULT_COLLECTION = "rag_collection"
DEFAULT_TOP_K = 4
PORTKEY_API_KEY = os.getenv("PORTKEY_API_KEY", "").strip()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY", "").strip()

# Utilities
def read_file_to_text(fp: str) -> str:
    name = fp.lower()
    with open(fp, "rb") as f:
        data = f.read()
    buf = io.BytesIO(data)
    if name.endswith(".pdf"):
        return "\n".join([p.extract_text() or "" for p in PdfReader(buf).pages])
    elif name.endswith(".docx"):
        return "\n".join([p.text for p in DocxDocument(buf).paragraphs])
    else:
        return data.decode("utf-8", errors="ignore")

def make_embeddings() -> HuggingFaceEmbeddings:
    return HuggingFaceEmbeddings(model_name=EMBED_MODEL_NAME)

def normalize_query(q: str) -> str:
    return q.lower().strip()

def qdrant_path() -> str:
    return os.getenv("QDRANT_PATH", "/content/qdrant").strip()

def get_qdrant_client() -> QdrantClient:
    path = qdrant_path()
    os.makedirs(path, exist_ok=True)
    return QdrantClient(path=path)

def ensure_collection(client: QdrantClient, name: str, vector_dim: int):
    names = [c.name for c in client.get_collections().collections or []]
    if name not in names:
        client.create_collection(
            collection_name=name,
            vectors_config=qmodels.VectorParams(size=vector_dim, distance=qmodels.Distance.COSINE),
        )

def count_vectors(client: QdrantClient, collection: str) -> int:
    try:
        return int(getattr(client.count(collection_name=collection, exact=True), "count", 0))
    except Exception:
        return 0

def upsert_chunks(client: QdrantClient, collection: str, chunks: List[str], embeddings: HuggingFaceEmbeddings):
    vectors = embeddings.embed_documents(chunks)
    points = [
        qmodels.PointStruct(id=str(uuid.uuid4()), vector=vec, payload={"text": text})
        for text, vec in zip(chunks, vectors)
    ]
    client.upsert(collection_name=collection, points=points, wait=True)

def build_index(client: QdrantClient, collection: str, text: str) -> int:
    splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=200)
    chunks = splitter.split_text(text)
    if not chunks:
        return 0
    embeddings = make_embeddings()
    dim = len(embeddings.embed_query("dimension probe"))
    ensure_collection(client, collection, vector_dim=dim)
    upsert_chunks(client, collection, chunks, embeddings)
    return len(chunks)

def clear_qdrant_storage_safe(client: QdrantClient) -> str:
    if client:
        try:
            client.close()
        except:
            pass
    path = qdrant_path()
    if os.path.exists(path):
        shutil.rmtree(path)
    os.makedirs(path, exist_ok=True)
    return "🧹 Qdrant storage cleared."

# Retriever
class QdrantRetriever:
    def __init__(self, client: QdrantClient, collection: str, embeddings: HuggingFaceEmbeddings, top_k: int = DEFAULT_TOP_K):
        self.client = client
        self.collection = collection
        self.embeddings = embeddings
        self.top_k = int(top_k)

    def invoke(self, query: str) -> List[Document]:
        qvec = self.embeddings.embed_query(normalize_query(query))
        result = self.client.query_points(
            collection_name=self.collection,
            query=qvec,
            limit=self.top_k,
            with_payload=True
        )
        points = result[0] if isinstance(result, tuple) else result

        def safe_extract_payload(point):
            if hasattr(point, "payload"):
                return point.payload
            if isinstance(point, tuple) and len(point) >= 2 and isinstance(point[1], dict):
                return point[1]
            return {}

        def safe_extract_score(point):
            return getattr(point, "score", 0.0)

        docs = []
        for p in points:
            payload = safe_extract_payload(p)
            text = payload.get("text", "")
            if text:
                docs.append(Document(page_content=text, metadata={"score": safe_extract_score(p)}))

        if not docs:
            scroll_result = self.client.scroll(collection_name=self.collection, limit=256, with_payload=True)
            all_points = scroll_result[0] if isinstance(scroll_result, tuple) else scroll_result
            for p in all_points:
                payload = safe_extract_payload(p)
                txt = payload.get("text", "")
                if normalize_query(query) in txt.lower():
                    docs.append(Document(page_content=txt, metadata={"score": 0.0}))
        return docs

# LLM Client
class LLMClient:
    def __init__(self):
        if not PORTKEY_API_KEY or not GOOGLE_API_KEY:
            raise RuntimeError("Missing API keys.")
        self.client = Portkey(api_key=PORTKEY_API_KEY, provider="google", Authorization=f"Bearer {GOOGLE_API_KEY}", timeout=30)

    def generate_text(self, prompt: str) -> str:
        resp = self.client.chat.completions.create(
            model="@<provide-your-int-portkey>/gemini-2.5-flash",
            messages=[{"role": "user", "content": prompt}],
            max_output_tokens=512
        )
        try:
            return (resp.choices[0].message.content or "").strip()
        except:
            return ""

# LangGraph pipeline
@dataclass
class RAGState:
    question: str
    context_docs: List[Document]
    answer: str

def make_graph(llm: LLMClient, allow_fallback: bool):
    def retrieve(state: RAGState): return state
    def generate(state: RAGState):
        context = "\n\n".join([d.page_content for d in state.context_docs]) if state.context_docs else ""
        prompt = f"Use ONLY the following context to answer:\n{context}\n\nQuestion:\n{state.question}\nAnswer:"
        ans = llm.generate_text(prompt)
        return RAGState(state.question, state.context_docs, ans)
    def fallback(state: RAGState):
        if not allow_fallback and not state.answer.strip():
            return RAGState(state.question, state.context_docs, "I don't know")
        low_score = True
        if state.context_docs:
            top_score = state.context_docs[0].metadata.get("score", 0.0)
            low_score = top_score < 0.3
        if allow_fallback and (not state.context_docs or low_score or not state.answer.strip()):
            return RAGState(state.question, [], llm.generate_text(state.question))
        return state
    g = StateGraph(RAGState)
    g.add_node("retrieve", retrieve)
    g.add_node("generate", generate)
    g.add_node("fallback", fallback)
    g.set_entry_point("retrieve")
    g.add_edge("retrieve", "generate")
    g.add_edge("generate", "fallback")
    g.add_edge("fallback", END)
    return g.compile()

# Sample doc
SAMPLE_DOC_TEXT = """Deployment Runbook:
1. Build container image.
2. Push to registry.
3. Apply manifest to cluster.
4. Validate health checks.
"""

# Gradio UI
with gr.Blocks(title="mAI") as demo:
    gr.Markdown("## mAI\n\n### Steps:\n1️⃣ Initialize\n2️⃣ Upload file & Build Index\n3️⃣ Query")
    state_client = gr.State(None)
    state_llm = gr.State(None)
    state_embeddings = gr.State(None)
    state_retriever = gr.State(None)
    boot_btn = gr.Button("Initialize")
    boot_status = gr.Markdown("")
    def on_boot():
        return get_qdrant_client(), LLMClient(), make_embeddings(), "✅ Initialized."
    boot_btn.click(on_boot, [], [state_client, state_llm, state_embeddings, boot_status])
    clear_storage_btn = gr.Button("Clear Qdrant Storage")
    def on_clear_storage(client):
        return clear_qdrant_storage_safe(client), None
    clear_storage_btn.click(on_clear_storage, [state_client], [boot_status, state_client])
    collection_in = gr.Textbox(label="Collection", value=DEFAULT_COLLECTION)
    file_in = gr.File(label="Upload file", file_types=[".pdf", ".docx", ".txt"], type="filepath")
    topk_in = gr.Slider(label="Top-K", minimum=1, maximum=10, step=1, value=DEFAULT_TOP_K)
    build_btn = gr.Button("Build Index")
    load_sample_btn = gr.Button("Load Sample Doc")
    build_status = gr.Markdown("")
    col_status = gr.Markdown("")
    debug_out = gr.Markdown("")
    def retriever_for(client, collection, topk, embeddings):
        return QdrantRetriever(client, collection, embeddings, top_k=topk)
    def on_build(fp, collection, topk, client, embeddings):
        if not client or not embeddings: return "❌ Initialize first.", "", "", None
        if not fp: return "❌ Upload a file.", "", "", None
        text = read_file_to_text(fp)
        n_chunks = build_index(client, collection, text)
        total = count_vectors(client, collection)
        return f"✅ Indexed {n_chunks} chunks.", f"Collection has {total} vectors.", f"Indexed {n_chunks}.", retriever_for(client, collection, topk, embeddings)
    build_btn.click(on_build, [file_in, collection_in, topk_in, state_client, state_embeddings], [build_status, col_status, debug_out, state_retriever])
    def on_load_sample(collection, topk, client, embeddings):
        if not client or not embeddings: return "❌ Initialize first.", "", "", None
        n_chunks = build_index(client, collection, SAMPLE_DOC_TEXT)
        total = count_vectors(client, collection)
        return f"✅ Sample doc indexed: {n_chunks} chunks.", f"Collection has {total} vectors.", f"Indexed sample.", retriever_for(client, collection, topk, embeddings)
    load_sample_btn.click(on_load_sample, [collection_in, topk_in, state_client, state_embeddings], [build_status, col_status, debug_out, state_retriever])
    allow_fallback_toggle = gr.Checkbox(label="Allow fallback to Gemini general knowledge?", value=True)
    query_in = gr.Textbox(label="Your query")
    submit_btn = gr.Button("Submit")
    answer_out = gr.Markdown("")
    sources_out = gr.Markdown("")
    status_query = gr.Markdown("")
    retrieval_debug = gr.Markdown("")
    def on_submit(query, retriever, llm, allow_fallback):
        try:
            debug_info = "### Debug Logs\n"
            if not query.strip():
                return "❌ Empty query.", "_No sources_", "⚠️", debug_info
            if not (retriever and llm):
                return "❌ Initialize and build index first.", "_No sources_", "⚠️", debug_info
            docs = retriever.invoke(query)
            retrieval_count = len(docs)
            debug_info += f"Retrieved {retrieval_count} docs.\n"
            if docs:
                debug_info += f"Top score: {docs[0].metadata.get('score', 0.0):.4f}\n"
            else:
                debug_info += "No docs retrieved.\n"
            graph = make_graph(llm, allow_fallback)
            final_dict = graph.invoke(RAGState(question=query, context_docs=docs, answer=""))
            final = RAGState(**final_dict)
            raw_answer = final.answer.strip() or "⚠️ LLM returned empty response."
            if final.context_docs:
                items = "\n\n".join([
                    f"**Context {i}**\n```\n{d.page_content[:500]}...\n```"
                    for i, d in enumerate(final.context_docs, 1)
                ])
                sources_md = f"### Context\n{items}"
            else:
                sources_md = "### Context\n_No context._"
            return f"### Answer\n\n{raw_answer}", sources_md, "✅ Done.", debug_info
        except AttributeError as e:
            methods = dir(retriever.client) if retriever else []
            return (
                f"❌ Error: {e}",
                "_No sources_",
                "⚠️",
                "### Debug Logs\n"
                + f"Exception: {str(e)}\n"
                + "Available methods:\n"
                + "\n".join([m for m in methods if any(k in m for k in ['query', 'search', 'scroll'])])
            )
        except Exception as e:
            return f"❌ Error: {e}", "_No sources_", "⚠️", f"### Debug Logs\nException: {str(e)}"
    submit_btn.click(on_submit, [query_in, state_retriever, state_llm, allow_fallback_toggle], [answer_out, sources_out, status_query, retrieval_debug])

demo.launch(share=True)
